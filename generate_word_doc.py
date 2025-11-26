from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime

STUDENT_KEY = "2511101ALLC"

def add_heading(doc: Document, text: str, level: int = 1):
	para = doc.add_heading(text, level=level)
	return para

def add_paragraph(doc: Document, text: str, size: int = 12):
	p = doc.add_paragraph(text)
	for run in p.runs:
		run.font.size = Pt(size)
	return p

def add_placeholder_image(doc: Document, title: str, instructions: str):
	doc.add_paragraph().add_run().add_break()
	doc.add_paragraph(f"{title}").bold = True
	add_paragraph(doc, instructions)
	add_paragraph(doc, "[Inserte captura aquí: clic en Insertar > Imágenes]", 10)

def build_doc():
	doc = Document()

	# Portada
	add_heading(doc, "Examen Final - Programación IX", level=0)
	add_paragraph(doc, "Nombre completo: ______________________________")
	add_paragraph(doc, "Carné: ______________________________")
	add_paragraph(doc, f"Fecha: {datetime.now().strftime('%Y-%m-%d')}")
	p = doc.add_paragraph("Clave única del estudiante: ")
	run = p.add_run(STUDENT_KEY)
	run.bold = True
	run.font.size = Pt(18)
	add_paragraph(doc, "Curso: Programación IX")
	add_paragraph(doc, "Docente: David Orozco")

	doc.add_page_break()

	# Sección 1: Explicación del Proyecto
	add_heading(doc, "Sección 1: Explicación del Proyecto", level=1)
	add_paragraph(doc, "Este proyecto implementa una API REST en Django/DRF para gestionar Proyectos Universitarios y un frontend en React (Vite) para consumirla.")
	add_paragraph(doc, "La API expone endpoints CRUD para el modelo Project, con campos name, student_name, description, status y timestamps. El listado permite filtrar por status mediante query params.")
	add_paragraph(doc, "Flujo: React realiza solicitudes HTTP (axios) a /api/projects/ para listar y crear proyectos. Django/DRF recibe la solicitud, la procesa mediante un ViewSet y responde con JSON. Tras crear un proyecto, React refresca la lista y la interfaz se actualiza mostrando el nuevo registro.")
	add_paragraph(doc, "Componentes: Project (models.py), ProjectSerializer (serializers.py), ProjectViewSet (views.py) y rutas con DefaultRouter en config/urls.py. En frontend: cliente HTTP (src/api/client.js), página ProjectsPage.jsx con lista, estados de carga/error y formulario de creación.")

	doc.add_page_break()

	# Sección 2: Capturas de Pantalla
	add_heading(doc, "Sección 2: Capturas de Pantalla", level=1)
	add_paragraph(doc, "Nota: Las capturas deben ser reales del proyecto del estudiante y deben mostrar la clave personal visible en la pantalla del frontend.")

	add_placeholder_image(doc, "2.1 Modelo Project (models.py)", "Muestre el archivo con la clase Project y sus campos.")
	add_placeholder_image(doc, "2.2 Endpoint /api/projects/", "Prueba en Postman o navegador mostrando la respuesta JSON.")
	add_placeholder_image(doc, "2.3 Filtro por status", "Ejemplo: /api/projects/?status=pending mostrando resultados filtrados.")
	add_placeholder_image(doc, "2.4 React con lista y clave visible", "Pantalla de ProjectsPage con la clave en la esquina y la lista cargada.")
	add_placeholder_image(doc, "2.5 Formulario de creación", "Pantalla mostrando el formulario antes de enviar.")
	add_placeholder_image(doc, "2.6 Después de crear", "Pantalla mostrando el proyecto nuevo agregado en la lista.")

	doc.add_page_break()

	# Sección 3: Comentario Final
	add_heading(doc, "Sección 3: Comentario Final", level=1)
	add_paragraph(doc, "¿Qué fue lo más difícil? ______________________________")
	add_paragraph(doc, "¿Qué aprendió? ______________________________")
	add_paragraph(doc, "¿Qué mejoraría? ______________________________")

	doc.save("Examen_Progra4_Documento.docx")

if __name__ == "__main__":
	build_doc()


